from zhconv import convert
from docx import Document
# zh-hant 简体转繁体
# zh-hans 繁体转简体
word = Document('a.docx')
rule = 'zh-hans'
for t in word.paragraphs:
    t.text = convert(t.text, rule)
for i in word.tables:    
    for p in i.rows:        
        for h in p.cells:            
            h.text = convert(h.text, rule)
word.save('a_jt.docx')